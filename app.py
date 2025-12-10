"""
LexBoost Bar 法考加速 - 完整版主程式
by AICowLaw法烤牛
"""

import streamlit as st
from datetime import datetime
import time
import random
import asyncio
import io
import base64

# ==================== 頁面配置 ====================
st.set_page_config(
    page_title="LexBoost Bar 法考加速",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==================== 效能優化配置 ====================
# 減少不必要的重新渲染
if 'initialized' not in st.session_state:
    st.session_state.initialized = True

# ==================== 科目列表 ====================
SUBJECTS = [
    "民法", "民訴法", "刑法", "刑訴法", "憲法", "行政法",
    "公司法", "證交法", "保險法", "票據法", "強執法",
    "國私法", "國公法", "法律倫理", "法學英文",
    "智財法", "海商海洋法", "勞社法", "財稅法", "其他"
]

# ==================== 自訂 CSS ====================
st.markdown("""
<style>
/* 隱藏 Streamlit 預設 */
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
header {visibility: hidden;}

/* ========== 整體配色 ========== */
.stApp {
    background: #f0f2f6;
}

[data-testid="stMainBlockContainer"] {
    background: white;
    border-radius: 20px;
    padding: 2rem;
    margin: 1rem;
    box-shadow: 0 4px 20px rgba(0,0,0,0.1);
}

[data-testid="stMainBlockContainer"] h1,
[data-testid="stMainBlockContainer"] h2,
[data-testid="stMainBlockContainer"] h3,
[data-testid="stMainBlockContainer"] h4,
[data-testid="stMainBlockContainer"] p,
[data-testid="stMainBlockContainer"] span,
[data-testid="stMainBlockContainer"] label,
[data-testid="stMainBlockContainer"] div {
    color: #1f2937 !important;
}

/* ========== 側邊欄 ========== */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #1f2937 0%, #111827 100%);
}

[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3,
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] div {
    color: white !important;
}

[data-testid="stSidebar"] .stButton > button {
    background: linear-gradient(135deg, #4b5563 0%, #374151 100%) !important;
    border: 1px solid rgba(255,255,255,0.2) !important;
    color: white !important;
}

[data-testid="stSidebar"] .stButton > button:hover {
    background: linear-gradient(135deg, #6b7280 0%, #4b5563 100%) !important;
}

[data-testid="stSidebar"] hr {
    border-color: rgba(255,255,255,0.2) !important;
}

/* ========== 按鈕 ========== */
.stButton > button {
    background: linear-gradient(135deg, #6b7280 0%, #4b5563 100%) !important;
    color: white !important;
    border: none !important;
    border-radius: 12px !important;
    padding: 0.75rem 2rem !important;
    font-weight: 600 !important;
    box-shadow: 0 4px 15px rgba(107, 114, 128, 0.4) !important;
}

.stButton > button:hover {
    transform: translateY(-2px);
    box-shadow: 0 8px 25px rgba(107, 114, 128, 0.5) !important;
}

/* ========== 輸入框 ========== */
.stTextInput > div > div > input,
.stTextArea > div > div > textarea {
    background: white !important;
    color: #1f2937 !important;
    border: 2px solid #e5e7eb !important;
    border-radius: 12px !important;
    opacity: 1 !important;
}

.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus {
    border-color: #6b7280 !important;
}

.stTextInput > label,
.stTextArea > label,
.stSelectbox > label {
    color: #1f2937 !important;
    font-weight: 600 !important;
}

.stSelectbox > div > div {
    background: white !important;
    color: #1f2937 !important;
}

/* ========== 統計卡片 ========== */
.stat-card {
    background: linear-gradient(135deg, #6b7280 0%, #4b5563 100%);
    border-radius: 16px;
    padding: 1.5rem;
    text-align: center;
    box-shadow: 0 8px 32px rgba(107, 114, 128, 0.3);
}

.stat-card .stat-label,
.stat-card .stat-number {
    color: white !important;
}

.stat-number {
    font-size: 2.5rem;
    font-weight: 800;
    margin: 0.5rem 0;
}

.stat-label {
    font-size: 0.9rem;
    opacity: 0.9;
}

/* ========== 標籤 ========== */
.tag {
    display: inline-block;
    background: linear-gradient(135deg, #6b7280 0%, #4b5563 100%);
    color: white !important;
    padding: 0.25rem 0.75rem;
    border-radius: 20px;
    font-size: 0.8rem;
    margin: 0.2rem;
    font-weight: 500;
}

/* ========== 訊息框 ========== */
.success-box {
    background: #d1fae5;
    border-left: 4px solid #10b981;
    padding: 1rem 1.5rem;
    border-radius: 12px;
    margin: 1rem 0;
    color: #065f46 !important;
}

.warning-box {
    background: #fef3c7;
    border-left: 4px solid #f59e0b;
    padding: 1rem 1.5rem;
    border-radius: 12px;
    margin: 1rem 0;
    color: #92400e !important;
}

.info-box {
    background: #dbeafe;
    border-left: 4px solid #3b82f6;
    padding: 1rem 1.5rem;
    border-radius: 12px;
    margin: 1rem 0;
    color: #1e40af !important;
}

.card {
    background: #f9fafb;
    border-radius: 16px;
    padding: 1.5rem;
    margin: 1rem 0;
    border: 1px solid #e5e7eb;
    color: #1f2937 !important;
}

/* ========== 對話氣泡 ========== */
.chat-bubble-user {
    background: linear-gradient(135deg, #6b7280 0%, #4b5563 100%);
    color: white !important;
    padding: 1rem 1.5rem;
    border-radius: 20px 20px 5px 20px;
    margin: 0.5rem 0;
    max-width: 80%;
    margin-left: auto;
}

.chat-bubble-ai {
    background: #f3f4f6;
    color: #1f2937 !important;
    padding: 1rem 1.5rem;
    border-radius: 20px 20px 20px 5px;
    margin: 0.5rem 0;
    max-width: 80%;
}

/* ========== 固定對話框容器 ========== */
.chat-container-fixed {
    position: sticky;
    top: 20px;
    z-index: 100;
}

/* ========== 進度條 ========== */
.progress-bar {
    background: linear-gradient(90deg, #6b7280 0%, #4b5563 100%);
    height: 100%;
    border-radius: 10px;
}

/* ========== 頁尾 ========== */
.footer {
    text-align: center;
    padding: 2rem;
    border-top: 1px solid #e5e7eb;
    margin-top: 3rem;
    color: #6b7280 !important;
}

.footer-brand {
    font-size: 1.2rem;
    font-weight: 700;
    background: linear-gradient(135deg, #6b7280 0%, #4b5563 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}

/* ========== Tab 樣式 ========== */
.stTabs [data-baseweb="tab-list"] {
    gap: 1rem;
    background: #f3f4f6;
    border-radius: 12px;
    padding: 0.5rem;
}

.stTabs [data-baseweb="tab"] {
    border-radius: 8px;
    padding: 0.75rem 1.5rem;
    font-weight: 600;
    color: #1f2937 !important;
}

.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, #6b7280 0%, #4b5563 100%);
    color: white !important;
}

.streamlit-expanderHeader {
    background: #f9fafb !important;
    color: #1f2937 !important;
    font-weight: 600 !important;
}

[data-testid="stMetricValue"] {
    color: #1f2937 !important;
}
</style>
""", unsafe_allow_html=True)

# ==================== 初始化 Session State ====================
if 'current_page' not in st.session_state:
    st.session_state.current_page = 'home'
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'chat_mode' not in st.session_state:
    st.session_state.chat_mode = 'reference'
if 'user_id' not in st.session_state:
    st.session_state.user_id = None
if 'delete_confirm' not in st.session_state:
    st.session_state.delete_confirm = {}
if 'chat_input_clear' not in st.session_state:
    st.session_state.chat_input_clear = 0
if 'note_tab' not in st.session_state:
    st.session_state.note_tab = 'text'  # 預設為文字輸入
if 'logging_in' not in st.session_state:
    st.session_state.logging_in = False
if 'show_loading' not in st.session_state:
    st.session_state.show_loading = False

# 快速登入，不顯示過渡頁面（避免殘影）
if st.session_state.logging_in:
    st.session_state.logging_in = False
    st.session_state.show_loading = False  # 清除載入提示
    st.rerun()

# 如果已登入，不顯示登入頁面
if st.session_state.user_id:
    # 跳過登入頁面，直接進入主程式
    pass

# ==================== 使用者登入 ====================
if not st.session_state.user_id:
    st.markdown("""
    <div style='text-align: center; padding: 3rem 0;'>
        <h1 style='font-size: 2.5rem; color: #1f2937;'>⚖️ LexBoost Bar 法考加速</h1>
        <p style='color: #6b7280; font-size: 1.2rem; margin-bottom: 2rem;'>請選擇使用者並輸入密碼登入</p>
    </div>
    """, unsafe_allow_html=True)
    
    USER_PASSWORDS = {
        "九水": "13134",
        "使用者A": "a",
        "使用者B": "b",
        "使用者C": "c",
        "使用者D": "d"
    }
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        user_name = st.selectbox("👤 選擇你的名字", list(USER_PASSWORDS.keys()))
        password = st.text_input("🔒 輸入密碼", type="password", placeholder="請輸入密碼")
        
        with st.expander("💡 查看密碼提示"):
            st.info("預設密碼:\n- 九水: 13134\n- 使用者A-D: a, b, c, d")
        
        # 顯示載入提示（如果有）
        if st.session_state.show_loading:
            st.info("⏳ 正在載入，請稍候...")
        
        if st.button("🚀 登入系統", use_container_width=True, type="primary"):
            if password == USER_PASSWORDS.get(user_name):
                # 使用 spinner 顯示載入訊息
                with st.spinner("⏳ 正在載入，請稍候..."):
                    st.session_state.user_id = user_name
                    st.session_state.logging_in = True
                    st.session_state.show_loading = False
                    time.sleep(0.3)  # 最小延遲讓用戶看到提示
                st.rerun()
            else:
                st.error("❌ 密碼錯誤！請重新輸入")
    
    # 確保登入頁面結束，避免殘影
    st.stop()

# ==================== 初始化系統 ====================
def init_system():
    """Silently initialize system without printing messages"""
    import sys
    from io import StringIO
    
    # 捕捉所有輸出
    old_stdout = sys.stdout
    sys.stdout = StringIO()
    
    try:
        from config import Config
        Config.validate()
        from ai_core import AICore
        from data_manager import DataManager
        result = AICore(), DataManager(), None
    except Exception as e:
        result = None, None, str(e)
    finally:
        # 恢復輸出
        sys.stdout = old_stdout
    
    return result

# 手動快取：只初始化一次
if 'ai_core' not in st.session_state or 'data_manager' not in st.session_state:
    ai_core, data_manager, init_error = init_system()
    st.session_state.ai_core = ai_core
    st.session_state.data_manager = data_manager
    st.session_state.init_error = init_error
else:
    ai_core = st.session_state.ai_core
    data_manager = st.session_state.data_manager
    init_error = st.session_state.init_error

# ==================== TTS 和下載輔助函數 ====================
async def generate_tts_audio(text: str, voice: str = "zh-CN-XiaoxiaoNeural", use_fallback: bool = True) -> bytes:
    """使用 Edge TTS 生成語音，失敗時自動切換到 Google TTS"""
    import re
    
    # 清理文字：移除所有可能導致問題的內容
    clean_text = text
    
    # 1. 移除程式碼區塊（包括 mermaid、python 等）
    clean_text = re.sub(r'```[\s\S]*?```', '', clean_text)
    
    # 2. 移除 HTML 標籤
    clean_text = re.sub(r'<[^>]+>', '', clean_text)
    
    # 3. 移除 Markdown 連結 [text](url)
    clean_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean_text)
    
    # 4. 移除 Markdown 格式符號
    clean_text = re.sub(r'[#*`_~\[\]{}|\\]', '', clean_text)
    
    # 5. 移除特殊符號，只保留中文、英文、數字和基本標點
    clean_text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s，。！？、；：「」『』（）\-\.\,\!\?\:\;]', '', clean_text)
    
    # 6. 移除多餘空白和換行
    clean_text = re.sub(r'\s+', ' ', clean_text)
    clean_text = clean_text.strip()
    
    # 7. 檢查是否有內容
    if not clean_text or len(clean_text) < 5:
        raise ValueError(f"文字內容清理後為空或過短（原始長度：{len(text)}，清理後：{len(clean_text)}）。可能包含過多特殊格式。")
    
    # 8. 限制長度
    if len(clean_text) > 2000:
        clean_text = clean_text[:2000] + "..."
    
    # 首先嘗試 Edge TTS（帶重試機制）
    edge_tts_error = None
    for attempt in range(2):  # 重試 2 次
        try:
            import edge_tts
            import asyncio
            
            # 設定超時時間
            communicate = edge_tts.Communicate(clean_text, voice)
            audio_data = b""
            chunk_count = 0
            
            # 使用 asyncio.wait_for 設定超時
            async def collect_audio():
                nonlocal audio_data, chunk_count
                async for chunk in communicate.stream():
                    if chunk["type"] == "audio":
                        audio_data += chunk["data"]
                        chunk_count += 1
            
            # 30 秒超時
            await asyncio.wait_for(collect_audio(), timeout=30.0)
            
            if audio_data and chunk_count > 0:
                return audio_data
            else:
                edge_tts_error = f"無法生成音訊（收到 {chunk_count} 個音訊片段）"
                
        except asyncio.TimeoutError:
            edge_tts_error = "Edge TTS 服務連線超時（30秒）"
            if attempt == 0:
                await asyncio.sleep(1)  # 重試前等待 1 秒
                continue
        except Exception as e:
            edge_tts_error = str(e)
            if attempt == 0 and "No audio was received" not in str(e):
                await asyncio.sleep(1)  # 重試前等待 1 秒
                continue
        
        break  # 如果是 "No audio was received" 錯誤，不重試
    
    # Edge TTS 失敗，嘗試使用 Google TTS 作為備用方案
    if use_fallback:
        try:
            from gtts import gTTS
            import io
            
            # 使用 gTTS 生成語音
            tts = gTTS(text=clean_text, lang='zh-TW', slow=False)
            
            # 將音訊儲存到記憶體
            audio_fp = io.BytesIO()
            tts.write_to_fp(audio_fp)
            audio_fp.seek(0)
            audio_data = audio_fp.read()
            
            if audio_data:
                # 成功使用 gTTS，但提示使用者
                import warnings
                warnings.warn(f"Edge TTS 失敗（{edge_tts_error}），已自動切換到 Google TTS")
                return audio_data
                
        except ImportError:
            # gTTS 未安裝
            raise ValueError(
                f"Edge TTS 服務無法使用：{edge_tts_error}\n\n"
                f"備用的 Google TTS 也未安裝。\n"
                f"請執行以下指令安裝：pip install gTTS\n\n"
                f"或者稍後再試 Edge TTS。"
            )
        except Exception as gtts_error:
            # gTTS 也失敗了
            raise ValueError(
                f"Edge TTS 失敗：{edge_tts_error}\n"
                f"Google TTS 也失敗：{gtts_error}\n\n"
                f"建議：\n"
                f"1. 檢查網路連線\n"
                f"2. 檢查防火牆設定\n"
                f"3. 稍後再試"
            )
    
    # 不使用備用方案，直接報錯
    raise ValueError(
        f"Edge TTS 服務無回應：{edge_tts_error}\n\n"
        f"可能原因：\n"
        f"1. 網路連線問題\n"
        f"2. Edge TTS 服務暫時無法使用\n"
        f"3. 防火牆阻擋\n\n"
        f"建議：請稍後再試，或檢查網路連線。\n\n"
        f"清理後文字預覽：{clean_text[:100]}..."
    )

def create_download_link(content: str, filename: str, file_format: str = "md") -> str:
    """創建下載連結"""
    if file_format == "md":
        b64 = base64.b64encode(content.encode()).decode()
        mime_type = "text/markdown"
    elif file_format == "txt":
        b64 = base64.b64encode(content.encode()).decode()
        mime_type = "text/plain"
    else:
        return None
    
    return f'<a href="data:{mime_type};base64,{b64}" download="{filename}.{file_format}">📥 下載 {file_format.upper()}</a>'

# ==================== 側邊欄 ====================
with st.sidebar:
    st.markdown(f"""
    <div style='text-align: center; padding: 1rem;'>
        <h1 style='color: white; font-size: 1.8rem; margin-bottom: 0;'>⚖️ LexBoost Bar</h1>
        <p style='color: #a5b4fc; font-size: 1rem;'>法考加速</p>
        <p style='color: #10b981; font-size: 0.9rem; margin-top: 0.5rem;'>👤 {st.session_state.user_id}</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    if data_manager:
        stats = data_manager.get_stats(st.session_state.user_id)
        st.markdown(f"""
        <div style='color: white; padding: 1rem;'>
            <div style='display: flex; justify-content: space-between; margin-bottom: 1rem;'>
                <span>📚 總筆記</span>
                <span style='font-weight: bold;'>{stats['total_notes']}</span>
            </div>
            <div style='display: flex; justify-content: space-between; margin-bottom: 1rem;'>
                <span>⏰ 待複習</span>
                <span style='font-weight: bold; color: #fbbf24;'>{stats['due_today']}</span>
            </div>
            <div style='display: flex; justify-content: space-between; margin-bottom: 1rem;'>
                <span>✅ 已複習</span>
                <span style='font-weight: bold; color: #34d399;'>{stats['reviewed']}</span>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    menu_items = [
        ("🏠", "首頁", "home"),
        ("📝", "建立筆記", "note"),
        ("💬", "AI 互動學習", "chat"),
        ("🔍", "智慧搜尋", "search"),
        ("🔄", "複習推薦", "review"),
        ("📚", "歷史資料庫", "database")
    ]
    
    for icon, label, page in menu_items:
        if st.button(f"{icon} {label}", key=f"nav_{page}", use_container_width=True):
            st.session_state.current_page = page
            st.rerun()
    
    st.markdown("---")
    
    if st.button("🚪 登出", use_container_width=True):
        st.session_state.user_id = None
        st.rerun()
    
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: #6b7280; font-size: 0.8rem; padding: 1rem;'>
        <p>🐄 AICowLaw法烤牛</p>
        <p>v1.0.0</p>
    </div>
    """, unsafe_allow_html=True)

# ==================== 錯誤處理 ====================
if init_error:
    st.error(f"⚠️ 系統初始化失敗：{init_error}")
    st.info("請檢查 .env 檔案中的 API 金鑰設定")
    st.stop()

# ==================== 頂部導航欄 ====================
def render_top_nav():
    st.markdown("""
    <h1 style='text-align: center; color: #1f2937; margin-bottom: 0.5rem;'>⚖️ LexBoost Bar 法考加速</h1>
    <p style='text-align: center; color: #6b7280; margin-bottom: 1.5rem;'>by AICowLaw法烤牛</p>
    """, unsafe_allow_html=True)
    
    cols = st.columns(6)
    pages = [("🏠 首頁", "home"), ("📝 建立筆記", "note"), ("💬 AI 互動學習", "chat"), 
             ("🔍 智慧搜尋", "search"), ("🔄 複習推薦", "review"), ("📚 歷史資料庫", "database")]
    
    for col, (label, page) in zip(cols, pages):
        with col:
            if st.button(label, key=f"top_{page}", use_container_width=True):
                st.session_state.current_page = page
                st.rerun()

# ==================== 首頁 ====================
def render_home():
    render_top_nav()
    
    st.markdown("""
    <div style='text-align: center; padding: 2rem 0;'>
        <h1 style='font-size: 2.5rem; color: #4b5563;'>
            🎯 讓 AI 成為你的法考加速器
        </h1>
        <p style='color: #6b7280; font-size: 1.2rem;'>科學化學習 × 智慧化複習 × 個人化推薦</p>
    </div>
    """, unsafe_allow_html=True)
    
    stats = data_manager.get_stats(st.session_state.user_id)
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown(f"""
        <div class="stat-card">
            <div class="stat-label">📚 總筆記數</div>
            <div class="stat-number">{stats['total_notes']}</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown(f"""
        <div class="stat-card">
            <div class="stat-label">⏰ 今日待複習</div>
            <div class="stat-number">{stats['due_today']}</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown(f"""
        <div class="stat-card">
            <div class="stat-label">✅ 已完成複習</div>
            <div class="stat-number">{stats['reviewed']}</div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    st.markdown("### 📅 今日複習推薦")
    
    due_notes = data_manager.get_due_notes(st.session_state.user_id)
    
    if due_notes:
        st.markdown(f'<div class="warning-box">⏰ 你有 {len(due_notes)} 則筆記需要複習！</div>', unsafe_allow_html=True)
        
        for i, note in enumerate(due_notes[:3]):
            with st.expander(f"📝 {note.get('title', '無標題')} - {note.get('category', '未分類')}", expanded=(i==0)):
                st.markdown(f"**複習次數**：{note.get('review_count', 0)} 次")
                st.markdown(f"**難度**：🎯 {note.get('difficulty', '中等')}")
                if st.button(f"🔄 立即複習", key=f"review_home_{note['id']}", use_container_width=True):
                    st.session_state.current_page = 'review'
                    st.rerun()
    else:
        st.markdown('<div class="success-box">🎉 目前沒有待複習的筆記！繼續保持！</div>', unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("### ⚡ 快速開始")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("📝 建立新筆記", use_container_width=True, type="primary"):
            st.session_state.current_page = 'note'
            st.rerun()
    with col2:
        if st.button("💬 開始 AI 對話", use_container_width=True):
            st.session_state.current_page = 'chat'
            st.rerun()
    with col3:
        if st.button("🔄 開始今日複習", use_container_width=True):
            st.session_state.current_page = 'review'
            st.rerun()

# ==================== AI 筆記生成 ====================
def render_note():
    render_top_nav()
    
    st.markdown("## 📝 AI 組織筆記建立")
    st.markdown('<div class="info-box">📌 輸入法條或筆記內容，AI 會自動整理成結構化筆記</div>', unsafe_allow_html=True)
    
    # 標籤選擇器（使用按鈕）
    st.markdown("### 📋 選擇輸入方式")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("✍️ 文字/文件輸入", use_container_width=True, 
                     type="primary" if st.session_state.note_tab == 'text' else "secondary",
                     key="tab_btn_text"):
            st.session_state.note_tab = 'text'
    
    with col2:
        if st.button("📷 圖片/PDF 辨識", use_container_width=True,
                     type="primary" if st.session_state.note_tab == 'ocr' else "secondary",
                     key="tab_btn_ocr"):
            st.session_state.note_tab = 'ocr'
    
    with col3:
        if st.button("🎙️ 語音輸入", use_container_width=True,
                     type="primary" if st.session_state.note_tab == 'voice' else "secondary",
                     key="tab_btn_voice"):
            st.session_state.note_tab = 'voice'
    
    st.markdown("---")
    
    # 根據選擇顯示對應內容
    if st.session_state.note_tab == 'text':
        # PDF 上傳選項
        st.markdown("#### 📄 上傳 PDF 文件（可選）")
        pdf_file = st.file_uploader("上傳 PDF", type=['pdf'], key="pdf_upload")
        
        if pdf_file:
            if st.button("📖 讀取 PDF 內容", use_container_width=True):
                try:
                    import fitz  # PyMuPDF
                    pdf_doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
                    pdf_text = ""
                    for page in pdf_doc:
                        pdf_text += page.get_text()
                    
                    st.session_state.pdf_content = pdf_text
                    st.success(f"✅ 已讀取 {pdf_doc.page_count} 頁內容")
                except Exception as e:
                    st.error(f"❌ 讀取失敗：{e}\n請先安裝：pip install PyMuPDF")
        
        st.markdown("---")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            title = st.text_input("📌 筆記標題", placeholder="例如：民法第184條 - 侵權行為")
        
        with col2:
            category = st.selectbox("📁 科目分類", SUBJECTS)
        
        # 筆記內容
        st.markdown("### 📝 筆記內容")
        content = st.text_area(
            "輸入你的筆記內容",
            height=300,
            value=st.session_state.get('pdf_content', ''),
            placeholder="在此輸入法律筆記內容...",
            key="main_content_input"
        )
        
        col1, col2, col3 = st.columns(3)
        with col1:
            note_type = st.selectbox("🎯 筆記類型", ["重點整理", "考點分析", "案例解析"])
        with col2:
            tags = st.text_input("🏷️ 標籤", placeholder="用逗號分隔，例如：侵權,重點")
        with col3:
            difficulty = st.selectbox(
                "🎯 難度標籤",
                ["極簡單", "簡單", "中等", "困難", "極困難"],
                index=2  # 預設「中等」
            )
        
        # AI 生成風格選擇
        st.markdown("### 🎨 AI 筆記風格設定")
        
        # 導入新的 Prompt 模板和風格管理器
        from prompt_templates import get_all_style_options, get_style_instruction
        # from custom_style_manager import CustomStyleManager  # 暂時停用，檔案已刪除
        
        # 初始化風格管理器（暂時停用）
        # if 'style_manager' not in st.session_state:
        #     st.session_state.style_manager = CustomStyleManager()
        # style_manager = st.session_state.style_manager
        
        # 載入風格選項（簡化版）
        style_presets = get_all_style_options()
        
        # 簡化的風格選擇介面
        selected_style = st.selectbox("選擇筆記風格", list(style_presets.keys()), index=0)
        
        # 處理風格指示
        if selected_style == "✏️ 自訂風格":
            custom_style = st.text_area("請描述你想要的筆記風格", height=100, 
                placeholder="例如：用條列式整理，每個重點不超過30字，加上記憶口訣",
                key="custom_style_input")
            style_instruction = get_style_instruction(selected_style, custom_style)
        else:
            style_instruction = get_style_instruction(
                selected_style,
                user_id=st.session_state.user_id,
                style_manager=style_manager
            )
            # 顯示風格說明
            with st.expander("👀 查看此風格說明"):
                preview_text = style_presets[selected_style]
                if len(preview_text) > 200:
                    st.info(preview_text[:200] + "...")
                else:
                    st.info(preview_text)
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        # AI 生成按鈕和清空按鈕
        col_gen, col_clear = st.columns([3, 1])
        
        with col_gen:
            if st.button("🤖 AI 生成筆記", use_container_width=True, type="primary", key="generate_btn"):
                if content:
                    with st.spinner("🔮 AI 正在整理筆記..."):
                        ai_notes = ai_core.generate_ai_notes(content, note_type, style_instruction)
                        st.session_state.generated_notes = ai_notes
                        st.session_state.note_metadata = {
                            'title': title or f"{category} - {note_type}",
                            'category': category,
                            'tags': [t.strip() for t in tags.split(",")] if tags else [],
                            'difficulty': difficulty  # 改為 difficulty
                        }
                        st.rerun()
                else:
                    st.warning("⚠️ 請輸入內容")
        
        with col_clear:
            if st.button("🗑️ 清空所有內容", use_container_width=True, key="clear_top"):
                # 清空所有相關的 session state
                if 'generated_notes' in st.session_state:
                    del st.session_state.generated_notes
                if 'note_metadata' in st.session_state:
                    del st.session_state.note_metadata
                if 'pdf_content' in st.session_state:
                    del st.session_state.pdf_content
                if 'mindmap_code' in st.session_state:
                    del st.session_state.mindmap_code
                if 'system_diagram' in st.session_state:
                    del st.session_state.system_diagram
                st.success("✅ 已清空所有內容！")
                time.sleep(0.5)
                st.rerun()
        
        # 顯示生成結果
        if 'generated_notes' in st.session_state and st.session_state.generated_notes:
            st.markdown('<div class="success-box">✅ 筆記生成完成！</div>', unsafe_allow_html=True)
            
            st.markdown("### 📋 生成結果")
            
            # 分兩欄：編輯區和預覽區
            col1, col2 = st.columns([1, 1])
            
            with col1:
                st.markdown("**編輯筆記內容**")
                edited_notes = st.text_area(
                    "編輯區",
                    value=st.session_state.generated_notes,
                    height=600,
                    key="edit_notes_area",
                    label_visibility="collapsed"
                )
                st.session_state.generated_notes = edited_notes
            
            with col2:
                st.markdown("**預覽效果**")
                # 使用 markdown 渲染預覽，添加固定高度容器
                st.markdown(f'<div style="height: 600px; overflow-y: auto; border: 1px solid #e5e7eb; border-radius: 8px; padding: 1rem;">{edited_notes}</div>', unsafe_allow_html=True)
            
            st.markdown("### 💾 儲存與下載")
            
            # 儲存選項
            col1, col2 = st.columns(2)
            with col1:
                save_to_db = st.checkbox("✅ 儲存到資料庫（Airtable 雲端同步）", value=True)
            with col2:
                add_to_kb = st.checkbox("✅ 加入知識庫（支援 AI 智慧搜尋）", value=True)
            
            # 按鈕區 - 改為 4 欄
            col1, col2, col3, col4 = st.columns([2, 2, 1, 1])
            
            with col1:
                if st.button("💾 儲存筆記", use_container_width=True, type="primary"):
                    meta = st.session_state.note_metadata
                    
                    if save_to_db:
                        data_manager.save_note(
                            user_id=st.session_state.user_id,
                            title=meta['title'],
                            content=st.session_state.generated_notes,
                            category=meta['category'],
                            tags=meta['tags'],
                            difficulty=meta['difficulty']  # 改為 difficulty
                        )
                    
                    if add_to_kb:
                        ai_core.add_to_knowledge_base(
                            content=st.session_state.generated_notes,
                            metadata={
                                'type': 'note',
                                'title': meta['title'],
                                'category': meta['category'],
                                'tags': meta['tags'],
                                'difficulty': meta['difficulty'],
                                'user_id': st.session_state.user_id,
                                'note_id': 'temp_' + str(int(time.time() * 1000))  # 暫時 ID，實際應該在儲存後使用真實 note_id
                            }
                        )
                    
                    st.success("✅ 儲存成功！")
                    # 不再自動清空，讓使用者自己決定
            
            with col2:
                # 下載格式選擇和按鈕
                download_format = st.selectbox("下載格式", ["Markdown", "Word", "PDF"], label_visibility="collapsed")
                
                if download_format == "Markdown":
                    st.download_button(
                        "⬇️ 下載筆記",
                        st.session_state.generated_notes,
                        f"{st.session_state.note_metadata['title']}.md",
                        "text/markdown",
                        use_container_width=True
                    )
                elif download_format == "Word":
                    try:
                        from docx import Document
                        from io import BytesIO
                        
                        doc = Document()
                        doc.add_heading(st.session_state.note_metadata['title'], 0)
                        for line in st.session_state.generated_notes.split('\n'):
                            if line.strip():
                                doc.add_paragraph(line)
                        
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        
                        st.download_button(
                            "⬇️ 下載筆記",
                            buffer,
                            f"{st.session_state.note_metadata['title']}.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    except ImportError:
                        st.error("請先安裝 python-docx：pip install python-docx")
                elif download_format == "PDF":
                    try:
                        from reportlab.lib.pagesizes import A4
                        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                        from reportlab.pdfbase import pdfmetrics
                        from reportlab.pdfbase.ttfonts import TTFont
                        from io import BytesIO
                        
                        buffer = BytesIO()
                        doc = SimpleDocTemplate(buffer, pagesize=A4)
                        story = []
                        
                        # 註冊中文字體（使用系統字體）
                        try:
                            pdfmetrics.registerFont(TTFont('Microsoft-JhengHei', 'C:/Windows/Fonts/msjh.ttc'))
                            font_name = 'Microsoft-JhengHei'
                        except:
                            font_name = 'Helvetica'
                        
                        styles = getSampleStyleSheet()
                        title_style = ParagraphStyle(
                            'CustomTitle',
                            parent=styles['Heading1'],
                            fontName=font_name,
                            fontSize=18,
                            spaceAfter=30,
                        )
                        body_style = ParagraphStyle(
                            'CustomBody',
                            parent=styles['BodyText'],
                            fontName=font_name,
                            fontSize=12,
                            leading=20,
                        )
                        
                        # 標題
                        story.append(Paragraph(st.session_state.note_metadata['title'], title_style))
                        story.append(Spacer(1, 12))
                        
                        # 內容
                        for line in st.session_state.generated_notes.split('\n'):
                            if line.strip():
                                story.append(Paragraph(line.replace('<', '&lt;').replace('>', '&gt;'), body_style))
                                story.append(Spacer(1, 6))
                        
                        doc.build(story)
                        buffer.seek(0)
                        
                        st.download_button(
                            "⬇️ 下載筆記",
                            buffer,
                            f"{st.session_state.note_metadata['title']}.pdf",
                            "application/pdf",
                            use_container_width=True
                        )
                    except ImportError:
                        st.error("請先安裝 reportlab：pip install reportlab")
            
            
            with col3:
                if st.button("🔄 重新生成", use_container_width=True):
                    # 保留 metadata，只清空生成的筆記
                    if 'generated_notes' in st.session_state:
                        del st.session_state.generated_notes
                    st.rerun()
            
            with col4:
                if st.button("🗑️ 清空所有", use_container_width=True, key="clear_bottom"):
                    # 清空所有相關的 session state
                    if 'generated_notes' in st.session_state:
                        del st.session_state.generated_notes
                    if 'note_metadata' in st.session_state:
                        del st.session_state.note_metadata
                    if 'pdf_content' in st.session_state:
                        del st.session_state.pdf_content
                    if 'mindmap_code' in st.session_state:
                        del st.session_state.mindmap_code
                    if 'system_diagram' in st.session_state:
                        del st.session_state.system_diagram
                    st.success("✅ 已清空所有內容！")
                    time.sleep(0.5)
                    st.rerun()
            
            # 心智圖生成
            st.markdown("---")
            st.markdown("### 🗺️ 視覺化工具")
            
            col1, col2 = st.columns(2)
            with col1:
                if st.button("🗺️ 生成心智圖（Mermaid）", use_container_width=True):
                    with st.spinner("🎨 生成中..."):
                        mindmap_code = ai_core.generate_mind_map(st.session_state.generated_notes)
                        st.session_state.mindmap_code = mindmap_code
                        st.rerun()
            
            with col2:
                if st.button("📊 生成法律體系圖", use_container_width=True):
                    with st.spinner("🎨 生成中..."):
                        system_diagram = ai_core.generate_legal_system_diagram(st.session_state.generated_notes)
                        st.session_state.system_diagram = system_diagram
                        st.rerun()
            
            # 顯示生成的筆記
            if 'ai_notes' in st.session_state:
                st.markdown("### ✅ AI 生成的筆記")
                st.markdown(st.session_state.ai_notes)
                
                # 儲存按鈕
                if st.button("💾 儲存筆記", use_container_width=True, type="primary", key="save_generated"):
                    data_manager.save_note(
                        user_id=st.session_state.user_id,
                        title=f"AI筆記 - {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                        content=st.session_state.ai_notes,
                        category=category,
                        tags=[t.strip() for t in tags.split(",")] if tags else [],
                        difficulty=difficulty
                    )
                    st.success("✅ 筆記已儲存！")
                    del st.session_state.ai_notes
                    st.rerun()
            
            # 顯示心智圖
            if 'mindmap_code' in st.session_state:
                st.markdown("**心智圖預覽**")
                
                mermaid_html = f"""
                <div class="mermaid">
                {st.session_state.mindmap_code}
                </div>
                <script type="module">
                  import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs';
                  mermaid.initialize({{ startOnLoad: true }});
                </script>
                """
                st.components.v1.html(mermaid_html, height=600, scrolling=True)
                
                with st.expander("📋 查看 Mermaid 程式碼"):
                    st.code(st.session_state.mindmap_code, language="mermaid")
            
            # 顯示體系圖
            if 'system_diagram' in st.session_state:
                st.markdown("**法律體系圖預覽**")
                
                diagram_html = f"""
                <div class="mermaid">
                {st.session_state.system_diagram}
                </div>
                <script type="module">
                  import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs';
                  mermaid.initialize({{ startOnLoad: true }});
                </script>
                """
                st.components.v1.html(diagram_html, height=600, scrolling=True)
                
                with st.expander("📋 查看 Mermaid 程式碼"):
                    st.code(st.session_state.system_diagram, language="mermaid")
        
        # 直接儲存筆記（不用 AI 生成）
        st.markdown("---")
        st.markdown("### ✍️ 或直接儲存原始筆記（不使用 AI）")
        
        if st.button("💾 直接儲存原始內容", use_container_width=True, key="save_raw"):
            if content:
                data_manager.save_note(
                    user_id=st.session_state.user_id,
                    title=title or f"{category} - 筆記",
                    content=content,
                    category=category,
                    tags=[t.strip() for t in tags.split(",")] if tags else [],
                    difficulty="中等"  # 直接儲存的筆記預設中等難度
                )
                st.success("✅ 筆記已儲存！")
            else:
                st.warning("⚠️ 請輸入內容")
    
    elif st.session_state.note_tab == 'ocr':
        st.markdown("### 📷 圖片/PDF 辨識")
        st.markdown('<div class="info-box">📸 上傳圖片或 PDF，AI 會自動辨識並完整呈現內容</div>', unsafe_allow_html=True)
        
        # 筆記設定（在上傳前）
        st.markdown("#### 📋 筆記設定")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            ocr_category = st.selectbox("📚 科目", SUBJECTS, key="ocr_cat")
        with col2:
            ocr_note_type = st.selectbox("🎯 筆記類型", ["重點整理", "考點分析", "案例解析"], key="ocr_type")
        with col3:
            ocr_tags = st.text_input("🏷️ 標籤", placeholder="用逗號分隔", key="ocr_tags")
        with col4:
            ocr_difficulty = st.selectbox(
                "🎯 難度",
                ["極簡單", "簡單", "中等", "困難", "極困難"],
                index=2,
                key="ocr_diff"
            )
        
        st.markdown("---")
        
        uploaded_file = st.file_uploader("選擇圖片或 PDF", type=['png', 'jpg', 'jpeg', 'pdf'], key="ocr_upload")
        
        if uploaded_file:
            col1, col2 = st.columns([1, 1])
            
            with col1:
                if uploaded_file.type != 'application/pdf':
                    st.image(uploaded_file, caption="上傳的圖片", use_column_width=True)
                else:
                    st.info(f"📄 已上傳 PDF：{uploaded_file.name}")
            
            with col2:
                if st.button("🔍 辨識並生成筆記", use_container_width=True, type="primary"):
                    with st.spinner("📸 辨識中..."):
                        try:
                            file_type = uploaded_file.type
                            
                            if 'pdf' in file_type:
                                # PDF 處理
                                import fitz
                                pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                                
                                full_text = ""
                                for page_num in range(pdf_document.page_count):
                                    page = pdf_document[page_num]
                                    full_text += page.get_text()
                                
                                ocr_prompt = f"""請整理以下 PDF 內容成完整的筆記，包括所有文字、圖表說明。

PDF 內容：
{full_text}

請完整呈現所有內容，包括圖表的文字說明。只輸出筆記內容。用繁體中文。"""
                                
                                response = ai_core.model.generate_content(ocr_prompt)
                                
                            else:
                                # 圖片處理
                                import PIL.Image
                                img = PIL.Image.open(uploaded_file)
                                
                                ocr_prompt = """請辨識圖片中的所有內容，包括文字、圖表、表格等，完整呈現成筆記。

要求：
1. 辨識所有文字內容
2. 如果有圖表，請描述圖表內容
3. 如果有表格，請用文字呈現表格資料
4. 保持原有的結構和層次

只輸出筆記內容。用繁體中文。"""
                                
                                response = ai_core.model.generate_content([ocr_prompt, img])
                            
                            # 儲存結果
                            st.session_state.ocr_result = response.text
                            st.session_state.ocr_metadata = {
                                'title': f"辨識 - {uploaded_file.name}",
                                'category': ocr_category,
                                'note_type': ocr_note_type,
                                'tags': [t.strip() for t in ocr_tags.split(",")] if ocr_tags else ['OCR'],
                                'difficulty': ocr_difficulty
                            }
                            st.rerun()
                            
                        except Exception as e:
                            st.error(f"❌ 辨識失敗：{e}")
                            st.info("請確認已安裝：pip install PyMuPDF Pillow")
        
        # 顯示辨識結果
        if 'ocr_result' in st.session_state:
            st.markdown("---")
            st.markdown("### ✅ 辨識結果")
            
            # AI 通順文字按鈕
            col1, col2 = st.columns([3, 1])
            with col1:
                if st.button("✨ AI 通順文字與排版", use_container_width=True, type="secondary"):
                    with st.spinner("🤖 AI 正在整理文字..."):
                        smoothed = ai_core.model.generate_content(
                            f"""請將以下 OCR 辨識的文字，整理成通順、完整、格式化的內容。

要求：
1. 修正錯字和不通順的地方
2. 補充必要的標點符號
3. 使用適當的段落和換行
4. 保留原有的法律術語和概念
5. 不要額外解釋或延伸內容
6. 使用 Markdown 格式美化排版（標題、列表、粗體等）

OCR 辨識文字：
{st.session_state.ocr_result}
"""
                        ).text
                        st.session_state.smoothed_ocr = smoothed
                        st.success("✅ 整理完成！")
                        st.rerun()
            
            with col2:
                if st.button("🗑️ 清空結果", use_container_width=True):
                    # 清空所有 OCR 相關的 session state
                    keys_to_delete = ['ocr_result', 'smoothed_ocr', 'ocr_metadata', 'ocr_upload']
                    for key in keys_to_delete:
                        if key in st.session_state:
                            del st.session_state[key]
                    st.success("✅ 已清空所有內容！")
                    st.rerun()  # 移除 sleep，直接 rerun
            
            # 顯示內容 - 改為兩欄對照，可編輯
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("#### 📄 原始辨識內容（可編輯）")
                edited_original = st.text_area(
                    "編輯原始內容",
                    value=st.session_state.ocr_result,
                    height=400,
                    key="edit_ocr_original",
                    label_visibility="collapsed"
                )
                # 更新 session state
                st.session_state.ocr_result = edited_original
                
                # 儲存原始筆記按鈕
                if st.button("💾 儲存原始筆記", use_container_width=True, type="secondary", key="save_ocr_original"):
                    meta = st.session_state.ocr_metadata
                    data_manager.save_note(
                        user_id=st.session_state.user_id,
                        title=meta['title'] + " (原始)",
                        content=edited_original,
                        category=meta['category'],
                        tags=meta['tags'],
                        difficulty="中等"
                    )
                    st.success("✅ 原始筆記已儲存！")
            
            with col2:
                if 'smoothed_ocr' in st.session_state:
                    st.markdown("#### ✨ AI 整理後的內容（可編輯）")
                    edited_ai = st.text_area(
                        "編輯 AI 內容",
                        value=st.session_state.smoothed_ocr,
                        height=400,
                        key="edit_ocr_ai",
                        label_visibility="collapsed"
                    )
                    # 更新 session state
                    st.session_state.smoothed_ocr = edited_ai
                    
                    # 儲存 AI 筆記按鈕
                    if st.button("💾 儲存 AI 筆記", use_container_width=True, type="primary", key="save_ocr_ai"):
                        meta = st.session_state.ocr_metadata
                        data_manager.save_note(
                            user_id=st.session_state.user_id,
                            title=meta['title'],
                            content=edited_ai,
                            category=meta['category'],
                            tags=meta['tags'],
                            difficulty="中等"
                        )
                        st.success("✅ AI 筆記已儲存！")
                else:
                    st.info("👈 點擊左側「AI 通順文字與排版」按鈕生成 AI 整理版本")
    
    elif st.session_state.note_tab == 'voice':
        st.markdown("### 🎙️ 語音輸入")
        st.markdown('<div class="info-box">🎤 上傳錄音檔，轉文字或生成筆記</div>', unsafe_allow_html=True)
        
        # 筆記設定
        st.markdown("#### 📋 筆記設定")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            voice_category = st.selectbox("📚 科目", SUBJECTS, key="voice_cat")
        with col2:
            voice_note_type = st.selectbox("🎯 筆記類型", ["重點整理", "考點分析", "案例解析"], key="voice_type")
        with col3:
            voice_tags = st.text_input("🏷️ 標籤", placeholder="用逗號分隔", key="voice_tags")
        with col4:
            voice_difficulty = st.selectbox(
                "🎯 難度",
                ["極簡單", "簡單", "中等", "困難", "極困難"],
                index=2,
                key="voice_diff"
            )
        
        st.markdown("---")
        
        audio_file = st.file_uploader("上傳音檔", type=['mp3', 'wav', 'm4a', 'ogg', 'webm'], key="audio_upload")
        
        if audio_file:
            st.audio(audio_file)
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("📝 原音轉文字", use_container_width=True, type="primary"):
                    # 使用 OpenAI Whisper API（雲端版本）
                    try:
                        from openai import OpenAI
                        import tempfile
                        import os
                        
                        # 初始化 OpenAI 客戶端
                        client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
                        
                        with st.spinner("🎙️ 正在轉錄音檔..."):
                            # 儲存上傳的音檔到臨時檔案
                            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(audio_file.name)[1]) as tmp_file:
                                tmp_file.write(audio_file.read())
                                tmp_path = tmp_file.name
                            
                            try:
                                # 使用 OpenAI Whisper API 轉錄
                                with open(tmp_path, "rb") as audio:
                                    transcript = client.audio.transcriptions.create(
                                        model="whisper-1",
                                        file=audio,
                                        language="zh"
                                    )
                                
                                st.markdown("### ✅ 轉錄結果")
                                st.markdown(transcript.text)
                                st.session_state.voice_transcription = transcript.text
                                st.caption("💰 費用：約 $0.006/分鐘")
                                
                            finally:
                                # 清理臨時檔案
                                if os.path.exists(tmp_path):
                                    os.unlink(tmp_path)
                    
                    except ImportError:
                        st.error("🚧 語音轉文字功能需要安裝 openai 套件")
                        st.code("pip install openai", language="bash")
                    except Exception as e:
                        st.error(f"❌ 轉錄失敗：{str(e)}")
                        if "OPENAI_API_KEY" in str(e) or "api_key" in str(e).lower():
                            st.warning("⚠️ 請設定 OPENAI_API_KEY 環境變數")
                        if "WinError 2" in str(e) or "ffmpeg" in str(e).lower():
                            st.warning("""
                            ⚠️ **ffmpeg 未安裝或未加入 PATH**
                            
                            請安裝 ffmpeg：
                            1. 下載：https://www.gyan.dev/ffmpeg/builds/
                            2. 解壓縮並將 bin 資料夾加入系統 PATH
                            3. 或使用 chocolatey: `choco install ffmpeg`
                            4. 重新啟動終端機和 Streamlit
                            """)
            
            with col2:
                if st.button("🤖 AI 整理筆記", use_container_width=True, type="primary"):
                    if 'voice_transcription' in st.session_state:
                        with st.spinner("🤖 AI 整理中..."):
                            # 使用預設風格
                            ai_notes = ai_core.generate_ai_notes(
                                st.session_state.voice_transcription, 
                                voice_note_type, 
                                "請用最精簡的方式整理，只保留核心要點"
                            )
                            st.session_state.voice_notes = ai_notes
                            st.success("✅ 整理完成！")
                            st.rerun()
                    else:
                        st.warning("⚠️ 請先進行「原音轉文字」")
            
            with col3:
                if st.button("🗑️ 清空結果", use_container_width=True, key="clear_voice_btn"):
                    # 清空所有語音相關的 session state
                    keys_to_delete = ['voice_transcription', 'voice_notes', 'audio_upload']
                    for key in keys_to_delete:
                        if key in st.session_state:
                            del st.session_state[key]
                    st.success("✅ 已清空所有內容！")
                    st.rerun()
            
            # 顯示轉錄結果
            if 'voice_transcription' in st.session_state:
                st.markdown("---")
                st.markdown("### 📝 轉錄結果")
                st.text_area("轉錄文字", st.session_state.voice_transcription, height=200, key="voice_trans_display")
            
            # 顯示整理後的筆記
            if 'voice_notes' in st.session_state:
                st.markdown("### ✅ 整理後的筆記")
                st.markdown(st.session_state.voice_notes)
                
                # 儲存按鈕
                if st.button("💾 儲存筆記", use_container_width=True, type="primary", key="save_voice"):
                    data_manager.save_note(
                        user_id=st.session_state.user_id,
                        title=f"語音筆記 - {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                        content=st.session_state.voice_notes,
                        category=voice_category,
                        tags=[t.strip() for t in voice_tags.split(",")] if voice_tags else ['語音'],
                        difficulty=voice_difficulty
                    )
                    st.success("✅ 筆記已儲存！")
                    del st.session_state.voice_transcription
                    del st.session_state.voice_notes
                    st.rerun()

# ==================== AI 互動學習 ====================
def render_chat():
    render_top_nav()
    
    st.markdown("## 💬 AI 互動學習")
    
    # 初始化 Quiz State
    if 'quiz_data' not in st.session_state:
        st.session_state.quiz_data = None
    if 'quiz_answered' not in st.session_state:
        st.session_state.quiz_answered = False
    
    # 模式選擇
    st.markdown("### 🎯 選擇對話模式")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📚 參考書模式", use_container_width=True, key="btn_ref", 
                    type="primary" if st.session_state.chat_mode == 'reference' else "secondary"):
            st.session_state.chat_mode = 'reference'
            st.session_state.chat_history = []
            st.rerun()
    with col2:
        if st.button("🎓 蘇格拉底問答", use_container_width=True, key="btn_soc",
                    type="primary" if st.session_state.chat_mode == 'socratic' else "secondary"):
            st.session_state.chat_mode = 'socratic'
            st.session_state.chat_history = []
            st.rerun()
    with col3:
        if st.button("🎮 爭點搶答遊戲", use_container_width=True, key="btn_game",
                    type="primary" if st.session_state.chat_mode == 'game' else "secondary"):
            st.session_state.chat_mode = 'game'
            st.session_state.chat_history = []
            st.session_state.quiz_data = None
            st.session_state.quiz_answered = False
            st.rerun()
    
    mode_desc = {
        'reference': '📚 **參考書模式**：直接回答你的法律問題，引用相關法條與判例（使用 RAG 搜尋知識庫）',
        'socratic': '🎓 **蘇格拉底問答**：用提問引導你思考，加深理解',
        'game': '🎮 **爭點搶答**：從你的筆記資料庫中出題，測試對法條與爭點的熟悉度'
    }
    st.markdown(f'<div class="info-box">{mode_desc[st.session_state.chat_mode]}</div>', unsafe_allow_html=True)
    
    # ==================== 爭點搶答模式 UI ====================
    if st.session_state.chat_mode == 'game':
        st.markdown("### 🎮 測驗開始")
        
        # 初始化選擇的科目和已出過的題目
        if 'quiz_selected_subject' not in st.session_state:
            st.session_state.quiz_selected_subject = "全部"
        if 'used_quiz_notes' not in st.session_state:
            st.session_state.used_quiz_notes = set()  # 記錄已使用的筆記 ID
        if 'quiz_start_time' not in st.session_state:
            st.session_state.quiz_start_time = None
        if 'game_started' not in st.session_state:
            st.session_state.game_started = False
        
        # 如果還沒開始遊戲，顯示開始畫面
        if not st.session_state.game_started:
            col1, col2 = st.columns([2, 1])
            with col1:
                quiz_subject = st.selectbox("📚 選擇出題科目", ["全部"] + SUBJECTS, key="quiz_subject_select")
            with col2:
                st.write("")  # 對齊
                st.write("")
                if st.button("🚀 開始測驗", type="primary", use_container_width=True):
                    st.session_state.game_started = True
                    st.session_state.quiz_selected_subject = quiz_subject
                    st.session_state.quiz_data = None
                    st.session_state.quiz_answered = False
                    st.session_state.used_quiz_notes = set()
                    st.rerun()
            
            # 顯示統計
            st.markdown("---")
            st.markdown("### 📊 你的筆記統計")
            all_notes = data_manager.get_all_notes(st.session_state.user_id)
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("總筆記數", len(all_notes))
            with col2:
                categories = set([n.get('category') for n in all_notes if n.get('category')])
                st.metric("涵蓋科目", len(categories))
            with col3:
                st.metric("可出題數", len(all_notes))
            
            # 停止執行，不繼續往下
            st.stop()
        
        else:
            # 遊戲進行中
            # 第一行：科目選擇器 + 確定按鈕
            col_subject, col_confirm = st.columns([3, 1])
            with col_subject:
                st.markdown("#### 📚 出題科目")
                quiz_subject_game = st.selectbox(
                    "選擇科目",
                    ["全部"] + SUBJECTS,
                    index=(["全部"] + SUBJECTS).index(st.session_state.quiz_selected_subject) if st.session_state.quiz_selected_subject in (["全部"] + SUBJECTS) else 0,
                    key="quiz_subject_game_select",
                    label_visibility="collapsed"
                )
            
            with col_confirm:
                st.write("")
                st.write("")
                if st.button("✅ 確定科目", use_container_width=True, type="primary", key="confirm_subject_btn"):
                    if quiz_subject_game != st.session_state.quiz_selected_subject:
                        st.session_state.quiz_selected_subject = quiz_subject_game
                        st.session_state.quiz_data = None
                        st.session_state.quiz_answered = False
                        st.success(f"✅ 已切換到「{quiz_subject_game}」科目")
                        st.rerun()
            
            # 第二行：控制按鈕（平行排列）+ 計數器顯示
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("🔄 換一題", type="secondary", use_container_width=True, key="change_question_btn"):
                    st.session_state.quiz_data = None
                    st.session_state.quiz_answered = False
                    st.session_state.quiz_start_time = None
                    st.rerun()
            
            with col2:
                if st.button("🔁 重置題庫", type="secondary", use_container_width=True, key="reset_quiz_btn"):
                    st.session_state.used_quiz_notes = set()
                    st.session_state.quiz_data = None
                    st.session_state.quiz_answered = False
                    st.success("✅ 題庫已重置！")
                    st.rerun()
            
            with col3:
                # 使用 st.empty() 強制更新顯示
                counter_placeholder = st.empty()
                current_count = len(st.session_state.used_quiz_notes)
                counter_placeholder.markdown(f"### 📊 已出題\n# {current_count}")
            
            st.markdown("---")
            
            # 如果還沒有題目，生成題目
            if not st.session_state.quiz_data:
                # 先更新計數器（在 spinner 之前）
                all_notes = data_manager.get_all_notes(st.session_state.user_id)
                if st.session_state.quiz_selected_subject != "全部":
                    filtered_notes = [n for n in all_notes if n.get('category') == st.session_state.quiz_selected_subject]
                else:
                    filtered_notes = all_notes
                
                available_notes = [n for n in filtered_notes if n['id'] not in st.session_state.used_quiz_notes]
                
                if not available_notes and filtered_notes:
                    st.warning("⚠️ 所有題目都已出過！請點擊『重置題庫』重新開始。")
                    st.stop()
                
                # 先決定要用哪個筆記並更新計數器
                note_content = None
                source_difficulty = "中等"
                selected_note = None
                
                # 決定出題來源：資料庫筆記 (70%) 或 AI隨機 (30%)
                if available_notes and random.random() > 0.3:
                    # 從資料庫選擇筆記出題
                    selected_note = random.choice(available_notes)
                    note_content = f"標題：{selected_note.get('title')}\n內容：{selected_note.get('content')}"
                    source_difficulty = selected_note.get('difficulty', '中等')
                    
                    # 更新計數器
                    print(f"\n=== [資料庫出題] 出題前計數: {len(st.session_state.used_quiz_notes)} ===")
                    st.session_state.used_quiz_notes.add(selected_note['id'])
                    print(f"=== 添加筆記 ID: {selected_note['id']} ===")
                    print(f"=== 出題後計數: {len(st.session_state.used_quiz_notes)} ===")
                    print(f"=== 當前所有ID: {st.session_state.used_quiz_notes} ===\n")
                else:
                    # AI隨機出題（當資料庫無筆記或 30% 機率）
                    random_id = f"random_{int(time.time() * 1000)}"
                    print(f"\n=== [AI隨機出題] 出題前計數: {len(st.session_state.used_quiz_notes)} ===")
                    st.session_state.used_quiz_notes.add(random_id)
                    print(f"=== 添加隨機 ID: {random_id} ===")
                    print(f"=== 出題後計數: {len(st.session_state.used_quiz_notes)} ===")
                    print(f"=== 當前所有ID: {st.session_state.used_quiz_notes} ===\n")
                
                # 生成題目
                with st.spinner("🧠 正在出題..."):
                    quiz_data = ai_core.generate_quiz_question(
                        content=note_content, 
                        category=st.session_state.quiz_selected_subject if st.session_state.quiz_selected_subject != "全部" else None
                    )
                    quiz_data['difficulty'] = source_difficulty
                    st.session_state.quiz_data = quiz_data
                    st.session_state.quiz_answered = False
                
                print(f"=== 題目生成完成，準備rerun ===")
                print(f"=== 最終計數: {len(st.session_state.used_quiz_notes)} ===\n")
                # 強制重新載入
                st.rerun()
            
            # 顯示題目（不再有倒數計時）
            if st.session_state.quiz_data:
                q_data = st.session_state.quiz_data
                
                # 顯示題目
                st.markdown(f"""
                <div style="background: white; padding: 2rem; border-radius: 12px; border-left: 5px solid #6b7280; box-shadow: 0 4px 6px rgba(0,0,0,0.1); margin-bottom: 1rem;">
                    <h3 style="color: #1f2937; margin-bottom: 0.5rem;">{q_data.get('question', '題目讀取錯誤')}</h3>
                    <div style="color: #6b7280; font-size: 0.9rem;">
                        <span>🎯 難度：{q_data.get('difficulty', '中等')}</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)
            else:
                # 已回答，不顯示倒數
                st.markdown(f"""
                <div style="background: white; padding: 2rem; border-radius: 12px; border-left: 5px solid #6b7280; box-shadow: 0 4px 6px rgba(0,0,0,0.1); margin-bottom: 1rem;">
                    <h3 style="color: #1f2937; margin-bottom: 0.5rem;">{q_data.get('question', '題目讀取錯誤')}</h3>
                    <div style="color: #6b7280; font-size: 0.9rem;">
                        <span>🎯 難度：{q_data.get('difficulty', '中等')}</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)
            
            # 選項區
            options = q_data.get('options', [])
            
            # 如果已經回答，顯示結果
            if st.session_state.quiz_answered:
                correct_idx = q_data.get('answer_index', 0)
                user_choice_idx = st.session_state.get('user_choice_idx', -1)
                
                for i, opt in enumerate(options):
                    if i == correct_idx:
                        # 正確答案一定要顯示為綠色
                        st.success(f"✅ {opt} (正確答案)")
                    elif i == user_choice_idx:
                        # 用戶選錯的答案顯示為紅色
                        st.error(f"❌ {opt} (你的選擇)")
                    else:
                        # 其他選項
                        st.info(f"⚪ {opt}")
                
                st.markdown("---")
                st.markdown("### 💡 解析")
                st.markdown(f"""
                <div style="background: #f3f4f6; padding: 1.5rem; border-radius: 10px;">
                    {q_data.get('explanation', '無解析')}
                </div>
                """, unsafe_allow_html=True)
                
                st.markdown("<br>", unsafe_allow_html=True)
                
                # 下一題按鈕 - 直接出下一題，不回選單
                if st.button("👉 下一題", type="primary", use_container_width=True):
                    # 重設狀態以觸發新題目生成
                    st.session_state.quiz_data = None
                    st.session_state.quiz_answered = False
                    st.rerun()
                    
            else:
                # 還沒回答，顯示按鈕
                for i, opt in enumerate(options):
                    if st.button(opt, key=f"quiz_opt_{i}", use_container_width=True):
                        st.session_state.quiz_answered = True
                        st.session_state.user_choice_idx = i
                        st.rerun()

    # ==================== 一般對話模式 (原有的聊天 UI) ====================
    else:
        # 對話區域
        st.markdown("### 💬 對話區")
        
        chat_container = st.container()
        
        with chat_container:
            if not st.session_state.chat_history:
                welcome = {
                    'reference': '你好！我是你的法律參考書助手 📚 有任何法律問題都可以問我！我會搜尋知識庫找最相關的內容回答你。',
                    'socratic': '讓我們用提問來探索法律概念 🎓 請告訴我你想討論什麼主題？',
                }
                if st.session_state.chat_mode in welcome:
                    st.markdown(f'<div class="chat-bubble-ai">🤖 {welcome[st.session_state.chat_mode]}</div>', unsafe_allow_html=True)
            
            for msg in st.session_state.chat_history:
                if msg['role'] == 'user':
                    st.markdown(f'<div class="chat-bubble-user">{msg["content"]}</div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="chat-bubble-ai">🤖 {msg["content"]}</div>', unsafe_allow_html=True)
        
        # 輸入區
        user_input = st.text_area("💬 輸入訊息...", height=120, 
                                  key=f"chat_input_area_{st.session_state.chat_input_clear}", 
                                  placeholder="輸入你的法律問題...")
        
        col1, col2 = st.columns([5, 1])
        with col2:
            send = st.button("📤 發送", use_container_width=True, type="primary")
        
        if send and user_input:
            st.session_state.chat_history.append({'role': 'user', 'content': user_input})
            
            with st.spinner("🤔 思考中..."):
                if st.session_state.chat_mode == 'reference':
                    search_results = ai_core.search_knowledge_base(user_input, top_k=3)
                    response = ai_core.answer_question_with_rag(user_input, search_results)
                elif st.session_state.chat_mode == 'socratic':
                    response = ai_core.chat_with_ai(user_input, st.session_state.chat_history, 'socratic')
                else:
                    response = "錯誤：請切換回正確的模式"
            
            st.session_state.chat_history.append({'role': 'assistant', 'content': response})
            st.session_state.chat_input_clear += 1
            st.rerun()
        
        # 控制按鈕
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🗑️ 清除對話", use_container_width=True):
                st.session_state.chat_history = []
                st.rerun()
        with col2:
            if st.session_state.chat_history:
                chat_export = "\n\n".join([f"{'我' if m['role']=='user' else 'AI'}: {m['content']}" for m in st.session_state.chat_history])
                st.download_button("📅 下載對話記錄", chat_export, "對話記錄.txt", use_container_width=True)

# ==================== 智慧搜尋 ====================
def render_search():
    render_top_nav()
    
    st.markdown("## 🔍 智慧搜尋系統")
    st.markdown('<div class="info-box">🧠 使用 RAG 語義搜尋，理解你的問題意圖，找到最相關的法律知識</div>', unsafe_allow_html=True)
    
    # 更大的搜尋框
    query = st.text_area("🔎 輸入問題或關鍵字", height=120, placeholder="例如：什麼情況會構成侵權行為？")
    
    col1, col2, col3 = st.columns([3, 2, 1])
    with col1:
        category_filter = st.selectbox("📁 篩選科目", ["全部"] + SUBJECTS)
    with col2:
        difficulty_filter = st.selectbox("🎯 難度篩選", ["全部", "極簡單", "簡單", "中等", "困難", "極困難"])
    with col3:
        max_results = st.number_input("結果數", 1, 10, 5)
    
    if st.button("🔍 開始搜尋", type="primary", use_container_width=True) and query:
        with st.spinner("🔎 搜尋中..."):
            cat = None if category_filter == "全部" else category_filter
            results = ai_core.search_knowledge_base(query, top_k=max_results, category=cat)
            
            # 根據難度篩選結果
            if difficulty_filter != "全部" and results:
                results = [r for r in results if r['metadata'].get('difficulty') == difficulty_filter]
        
        if results:
            st.success(f"✅ 找到 {len(results)} 個相關結果")
            
            for i, r in enumerate(results):
                with st.expander(f"📝 {r['metadata'].get('title', f'結果 {i+1}')} - {r['metadata'].get('category', '未分類')} ({r['score']:.0%} 相關)", expanded=(i==0)):
                    st.markdown(f"**相關度**：{r['score']:.0%}")
                    st.markdown(f"**分類**：{r['metadata'].get('category', '未分類')}")
                    st.markdown(f"**難度**：🎯 {r['metadata'].get('difficulty', '中等')}")
                    if r['metadata'].get('tags'):
                        tags_str = " ".join([f"`{tag}`" for tag in r['metadata'].get('tags', []) if tag])
                        if tags_str.strip():
                            st.markdown(f"**標籤**：{tags_str}")
                    st.markdown("---")
                    st.markdown(r['content'])
        else:
            st.warning("😕 沒有找到相關結果，請嘗試其他關鍵字")

# ==================== 複習推薦 ====================
def render_review():
    render_top_nav()
    
    st.markdown("## 🔄 智慧複習推薦")
    
    # 每次都重新取得待複習筆記
    due_notes = data_manager.get_due_notes(st.session_state.user_id)
    
    # 簡化統計
    col1, col2 = st.columns(2)
    with col1:
        st.metric("📚 待複習筆記", len(due_notes))
    with col2:
        st.metric("✅ 今日已複習", data_manager.get_stats(st.session_state.user_id)['reviewed'])
    
    if due_notes:
        st.markdown(f'<div class="warning-box">⏰ 你有 {len(due_notes)} 則筆記需要複習！開始吧！</div>', unsafe_allow_html=True)
        
        note = due_notes[0]
        
        st.markdown(f"""
        <div class="card">
            <h3>{note.get('title', '無標題')}</h3>
            <div style="margin: 1rem 0;">
                <span class="tag">{note.get('category', '未分類')}</span>
                <span style="color: #6b7280; margin-left: 1rem;">📝 已複習 {note.get('review_count', 0)} 次</span>
                <span style="color: #6b7280; margin-left: 1rem;">🎯 難度: {note.get('difficulty', '中等')}</span>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        with st.expander("📖 查看內容", expanded=True):
            st.markdown(note.get('content', '無內容'))
        
        st.markdown("### 💭 記憶程度")
        st.markdown("""
        <div style="font-size: 0.85rem; color: #6b7280; margin-bottom: 0.5rem;">
        根據你的記憶程度安排下次複習：<br>
        完全不記得(1天) / 有點印象(3天) / 大致記得(7天) / 很熟悉(14天) / 完全精通(30天)
        </div>
        """, unsafe_allow_html=True)
        
        # 記住上次的選擇
        if 'last_memory_level' not in st.session_state:
            st.session_state.last_memory_level = 2  # 預設「大致記得」
        
        memory_level = st.radio(
            "選擇",
            ["❌ 完全不記得", "😐 有點印象", "😊 大致記得", "✅ 很熟悉", "🌟 完全精通"],
            horizontal=False,
            index=st.session_state.last_memory_level,
            key=f"memory_{note['id']}",
            label_visibility="collapsed"
        )
        
        if st.button("✅ 確認並進入下一則", use_container_width=True, type="primary"):
            level_map = {
                "❌ 完全不記得": "再次",
                "😐 有點印象": "困難", 
                "😊 大致記得": "良好",
                "✅ 很熟悉": "容易",
                "🌟 完全精通": "精通"
            }
            
            # 記住這次的選擇
            memory_options = ["❌ 完全不記得", "😐 有點印象", "😊 大致記得", "✅ 很熟悉", "🌟 完全精通"]
            st.session_state.last_memory_level = memory_options.index(memory_level)
            
            data_manager.update_review_schedule(note['id'], level_map[memory_level], st.session_state.user_id)
            st.rerun()
    else:
        st.markdown('<div class="success-box">🎉 太棒了！目前沒有待複習的筆記！</div>', unsafe_allow_html=True)
        st.balloons()

# ==================== 歷史資料庫 ====================
def render_database():
    render_top_nav()
    
    st.markdown("## 📚 統一歷史資料庫")
    
    # 篩選
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        filter_cat = st.selectbox("📁 科目篩選", ["全部"] + SUBJECTS)
    with col2:
        filter_difficulty = st.selectbox("🎯 難度篩選", ["全部", "極簡單", "簡單", "中等", "困難", "極困難"])
    with col3:
        filter_keyword = st.text_input("🔍 關鍵字")
    with col4:
        sort_by = st.selectbox("排序", ["最新優先", "最舊優先", "標題"])
    
    notes = data_manager.get_all_notes(st.session_state.user_id)
    
    # 篩選
    if filter_cat != "全部":
        notes = [n for n in notes if n.get('category') == filter_cat]
    if filter_difficulty != "全部":
        notes = [n for n in notes if n.get('difficulty') == filter_difficulty]
    if filter_keyword:
        notes = [n for n in notes if filter_keyword.lower() in n.get('title', '').lower() or 
                                     filter_keyword.lower() in n.get('content', '').lower()]
    
    # 排序
    if sort_by == "最新優先":
        notes = sorted(notes, key=lambda x: x.get('created_at', ''), reverse=True)
    elif sort_by == "最舊優先":
        notes = sorted(notes, key=lambda x: x.get('created_at', ''))
    else:
        notes = sorted(notes, key=lambda x: x.get('title', ''))
    
    st.markdown(f"### 📝 共 {len(notes)} 則筆記")
    
    if notes:
        for note in notes:
            with st.expander(f"📝 {note.get('title', '無標題')} - {note.get('category', '未分類')}", expanded=False):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(f"**建立時間**：{note.get('created_at', 'N/A')[:10]}")
                    st.markdown(f"**複習次數**：{note.get('review_count', 0)} 次")
                    st.markdown(f"**難度**：🎯 {note.get('difficulty', '中等')}")
                    if note.get('tags'):
                        tags_display = " ".join([f"`{tag}`" for tag in note.get('tags', []) if tag])
                        if tags_display.strip():
                            st.markdown(f"**標籤**：{tags_display}")
                
                with col2:
                    # TTS 按鈕
                    if st.button("🔊 轉語音", key=f"tts_{note['id']}", use_container_width=True):
                        with st.spinner("🎙️ 正在生成語音..."):
                            try:
                                import warnings
                                
                                # 捕獲警告
                                with warnings.catch_warnings(record=True) as w:
                                    warnings.simplefilter("always")
                                    
                                    # 生成 TTS
                                    audio_bytes = asyncio.run(generate_tts_audio(note.get('content', '')))
                                    
                                    # 檢查是否有警告（表示使用了備用 TTS）
                                    if w:
                                        for warning in w:
                                            if "Edge TTS 失敗" in str(warning.message):
                                                st.info("ℹ️ Edge TTS 暫時無法使用，已自動切換到 Google TTS")
                                
                                # 提供播放和下載
                                st.success("✅ 語音生成成功！")
                                st.audio(audio_bytes, format='audio/mp3')
                                st.download_button(
                                    label="📥 下載語音",
                                    data=audio_bytes,
                                    file_name=f"{note.get('title', 'note')}.mp3",
                                    mime="audio/mp3",
                                    key=f"download_audio_{note['id']}"
                                )
                            except Exception as e:
                                st.error(f"❌ 語音生成失敗：{e}")
                    
                    # 下載筆記按鈕 - 格式選擇
                    download_format = st.selectbox(
                        "下載格式",
                        ["Markdown (.md)", "Word (.docx)"],
                        key=f"format_{note['id']}",
                        label_visibility="collapsed"
                    )
                    
                    note_content = f"# {note.get('title', '無標題')}\n\n{note.get('content', '')}"
                    
                    if "Word" in download_format:
                        # Word 格式
                        try:
                            from docx import Document
                            from docx.shared import Pt, RGBColor
                            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                            import io
                            
                            doc = Document()
                            
                            # 標題
                            title = doc.add_heading(note.get('title', '無標題'), 0)
                            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            
                            # 元資料
                            meta = doc.add_paragraph()
                            meta.add_run(f"科目: {note.get('category', '未分類')}  |  ").bold = True
                            meta.add_run(f"難度: {note.get('difficulty', '中等')}  |  ").bold = True
                            meta.add_run(f"複習次數: {note.get('review_count', 0)}次").bold = True
                            meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            
                            doc.add_paragraph()  # 空行
                            
                            # 內容 - 處理 Markdown 格式
                            content_lines = note.get('content', '').split('\n')
                            for line in content_lines:
                                if line.startswith('# '):
                                    doc.add_heading(line[2:], level=1)
                                elif line.startswith('## '):
                                    doc.add_heading(line[3:], level=2)
                                elif line.startswith('### '):
                                    doc.add_heading(line[4:], level=3)
                                elif line.strip():
                                    doc.add_paragraph(line)
                            
                            # 儲存到記憶體
                            docx_io = io.BytesIO()
                            doc.save(docx_io)
                            docx_io.seek(0)
                            
                            st.download_button(
                                label="📥 下載 Word",
                                data=docx_io.getvalue(),
                                file_name=f"{note.get('title', 'note')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_note_{note['id']}",
                                use_container_width=True
                            )
                        except ImportError:
                            st.error("❌ Word 格式需要安裝 python-docx\n請執行: pip install python-docx")
                            st.download_button(
                                label="📥 下載 Markdown",
                                data=note_content,
                                file_name=f"{note.get('title', 'note')}.md",
                                mime="text/markdown",
                                key=f"download_note_{note['id']}",
                                use_container_width=True
                            )
                    else:
                        # Markdown 格式
                        st.download_button(
                            label="📥 下載 Markdown",
                            data=note_content,
                            file_name=f"{note.get('title', 'note')}.md",
                            mime="text/markdown",
                            key=f"download_note_{note['id']}",
                            use_container_width=True
                        )
                    
                    # 編輯按鈕
                    edit_key = f"edit_{note['id']}"
                    if st.button("✏️ 編輯", key=edit_key, use_container_width=True):
                        st.session_state.editing_note = note['id']
                        st.rerun()
                    
                    # 刪除按鈕
                    delete_key = f"del_{note['id']}"
                    if st.button("🗑️ 刪除", key=delete_key, type="secondary", use_container_width=True):
                        st.session_state.delete_confirm[note['id']] = True
                        st.rerun()
                    
                    if st.session_state.delete_confirm.get(note['id'], False):
                        st.warning("⚠️ 確定要刪除嗎？")
                        col_yes, col_no = st.columns(2)
                        with col_yes:
                            if st.button("✅ 確定", key=f"confirm_{note['id']}"):
                                data_manager.delete_note(note['id'], st.session_state.user_id)
                                ai_core.delete_from_knowledge_base(note['id'])
                                st.session_state.delete_confirm[note['id']] = False
                                st.success("✅ 已刪除")
                                st.rerun()
                        with col_no:
                            if st.button("❌ 取消", key=f"cancel_{note['id']}"):
                                st.session_state.delete_confirm[note['id']] = False
                                st.rerun()
                
                st.markdown("---")
                
                # 編輯模式
                if st.session_state.get('editing_note') == note['id']:
                    st.markdown("### ✏️ 編輯筆記")
                    
                    new_title = st.text_input("標題", value=note.get('title', ''), key=f"edit_title_{note['id']}")
                    new_content = st.text_area("內容", value=note.get('content', ''), height=300, key=f"edit_content_{note['id']}")
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        new_category = st.selectbox("科目", SUBJECTS, index=SUBJECTS.index(note.get('category', '民法')) if note.get('category') in SUBJECTS else 0, key=f"edit_cat_{note['id']}")
                    with col2:
                        new_difficulty = st.selectbox("難度", ["極簡單", "簡單", "中等", "困難", "極困難"], 
                                                     index=["極簡單", "簡單", "中等", "困難", "極困難"].index(note.get('difficulty', '中等')) if note.get('difficulty') in ["極簡單", "簡單", "中等", "困難", "極困難"] else 2,
                                                     key=f"edit_diff_{note['id']}")
                    with col3:
                        tags_str = ','.join(note.get('tags', []))
                        new_tags = st.text_input("標籤", value=tags_str, key=f"edit_tags_{note['id']}")
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("💾 儲存修改", use_container_width=True, type="primary", key=f"save_edit_{note['id']}"):
                            # 更新筆記
                            update_data = {
                                'title': new_title,
                                'content': new_content,
                                'category': new_category,
                                'difficulty': new_difficulty,
                                'tags': ','.join([t.strip() for t in new_tags.split(',')]) if new_tags else ''
                            }
                            data_manager.update_note(note['id'], update_data)
                            st.session_state.editing_note = None
                            st.success("✅ 修改已儲存！")
                            st.rerun()
                    with col2:
                        if st.button("❌ 取消編輯", use_container_width=True, key=f"cancel_edit_{note['id']}"):
                            st.session_state.editing_note = None
                            st.rerun()
                else:
                    # 顯示模式
                    st.markdown(note.get('content', '無內容'))
        
        # 匯出
        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            all_content = "\n\n---\n\n".join([f"# {n.get('title')}\n\n{n.get('content', '')}" for n in notes])
            st.download_button("📥 匯出所有筆記 (Markdown)", all_content, "所有筆記.md", use_container_width=True)
    else:
        st.info("📝 目前沒有筆記，開始建立你的第一則筆記吧！")

# ==================== 系統設定 ====================
def render_settings():
    render_top_nav()
    
    st.markdown("## ⚙️ 系統設定")
    
    st.markdown("### 📊 知識庫狀態")
    try:
        stats = ai_core.get_index_stats()
        col1, col2 = st.columns(2)
        with col1:
            st.metric("向量總數", stats['total_vectors'])
        with col2:
            st.metric("索引維度", stats['dimension'])
        st.success("✅ Pinecone 知識庫運作正常")
    except Exception as e:
        st.error(f"❌ 無法連接知識庫：{e}")
    
    st.markdown("### 🗃️ Airtable 狀態")
    try:
        notes = data_manager.get_all_notes(st.session_state.user_id)
        st.metric("你的筆記數", len(notes))
        st.success("✅ Airtable 連接正常")
    except Exception as e:
        st.error(f"❌ 無法連接 Airtable：{e}")
    
    st.markdown("### 🗑️ 資料管理")
    
    st.warning("⚠️ 以下操作不可逆，請謹慎操作！")
    
    st.markdown("#### 清空所有筆記")
    delete_password = st.text_input("請輸入密碼以確認刪除", type="password", key="delete_pwd")
    
    if st.button("🗑️ 清空所有筆記", type="secondary"):
        if delete_password == "delete":
            notes = data_manager.get_all_notes(st.session_state.user_id)
            for note in notes:
                data_manager.delete_note(note['id'], st.session_state.user_id)
            st.success("✅ 已清空所有筆記")
            time.sleep(1)
            st.rerun()
        else:
            st.error("❌ 密碼錯誤！無法刪除")
    
    st.caption("💡 提示：刪除密碼為 'delete'")
    
    st.markdown("### ℹ️ 關於")
    st.markdown("""
    - **版本**：v1.0.0
    - **開發團隊**：AICowLaw法烤牛
    - **技術棧**：Streamlit + Gemini + Pinecone + Airtable
    - **使用者系統**：支援多人使用，資料獨立
    - **功能**：
      - ✅ AI 筆記生成（7種風格）
      - ✅ OCR 圖片/PDF 辨識
      - ✅ 智慧搜尋（RAG）
      - ✅ SuperMemo 複習排程
      - ✅ 心智圖/體系圖生成
      - ✅ AI 互動學習（3種模式）
      - ✅ 語音轉文字（需安裝 Whisper）
    """)
    
    st.markdown("### 🔧 測試工具")
    
    with st.expander("🧪 測試 Pinecone 向量搜尋"):
        st.markdown("""
        **如何測試：**
        1. 建立一則筆記並勾選「加入知識庫」
        2. 到「智慧搜尋」輸入相關問題
        3. 應該會找到你的筆記
        """)
        
        if st.button("🔍 測試搜尋「侵權行為」"):
            results = ai_core.search_knowledge_base("侵權行為", top_k=3)
            if results:
                st.success(f"✅ 找到 {len(results)} 筆相關結果")
                for r in results:
                    st.write(f"- {r['metadata'].get('title')} ({r['score']:.0%})")
            else:
                st.warning("❌ 沒有找到結果，請先建立筆記並加入知識庫")
    
    with st.expander("🧪 測試 SuperMemo 複習排程"):
        st.markdown("""
        **如何測試：**
        1. 建立一則筆記
        2. 到「複習推薦」選擇記憶程度
        3. 到 Airtable 查看 `next_review` 欄位
        4. 應該會根據你的選擇計算下次複習時間
        """)
        
        st.info("💡 複習間隔：完全不記得(1天) / 有點印象(3天) / 大致記得(7天) / 很熟悉(14天) / 完全精通(30天)")

# ==================== 頁尾 ====================
def render_footer():
    st.markdown("""
    <div class="footer">
        <p class="footer-brand">⚖️ LexBoost Bar 法考加速</p>
        <p>by 🐄 AICowLaw法烤牛</p>
        <p style="font-size: 0.8rem; margin-top: 1rem;">讓 AI 成為你的法考加速器 🚀</p>
    </div>
    """, unsafe_allow_html=True)

# ==================== 主程式 ====================
def main():
    page = st.session_state.current_page
    
    if page == 'home':
        render_home()
    elif page == 'note':
        render_note()
    elif page == 'chat':
        render_chat()
    elif page == 'search':
        render_search()
    elif page == 'review':
        render_review()
    elif page == 'database':
        render_database()
    
    render_footer()

if __name__ == "__main__":
    main()